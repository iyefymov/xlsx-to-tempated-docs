#!/usr/bin/env python3
"""
Script to generate individual Word documents from Excel data using a template.

Usage:
    python generate_documents.py [--inspect] [--dry-run] [--pdf] [--pdf-only]

Options:
    --inspect    Only show Excel columns and template placeholders, don't generate
    --dry-run    Show what would be generated without creating files
    --pdf        Convert generated .docx files to PDF after creation
    --pdf-only   Convert existing .docx files in output dir to PDF (no generation)
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import unicodedata
from pathlib import Path

import pandas as pd
from docx import Document

# Configuration
SCRIPT_DIR = Path(__file__).parent
EXCEL_FILE = SCRIPT_DIR / "dataset.xlsx"
EXCEL_SHEET = "2. Filtered-Complete EOIs"  # The sheet with actual data
TEMPLATE_FILE = SCRIPT_DIR / "template.docx"
OUTPUT_DIR = SCRIPT_DIR / "output"
OUTPUT_PDF_DIR = SCRIPT_DIR / "output_pdf"

# Mapping from template placeholders to Excel column names
# Template uses «Placeholder_Name» format (mail merge style)
PLACEHOLDER_TO_COLUMN = {
    "PI_Name_Ucalgary_System": "PI Name (Ucalgary System)",
    "PI_Faculty_Ucalgary_system": "PI Faculty (Ucalgary system)",
    "PI_Department": "PI Department",
    "PI_Department_Ucalgary_system": "PI Department (Ucalgary system)",
    "Training_Environment": "Training Environment",
    "Project_Title": "Project Title",
    "Nomination_Type": "Nomination Type",
    "Applicable_Impact_Priority_Areas": "Applicable Impact+ Priority Areas",
    "Applicable_UCalgary_Transdisciplinary_Ar": "Applicable UCalgary Transdisciplinary Areas of Focus",
    "Most_Relevant_TriAgency": "Most Relevant Tri-Agency",
    "Abstract": "Abstract",
    "International_PostDoc_Count_2024": "International Post-Doc Count (2024)",
    "National_PostDoc_Count_2024": "National Post-Doc Count (2024)",
    "Grad_Student_Count_2024": "Grad Student Count (2024)",
}


def normalize_filename(text: str) -> str:
    """
    Normalize text for use in filenames.
    - Converts accented characters to ASCII equivalents (ò -> o, é -> e, etc.)
    - Removes or replaces invalid filename characters
    """
    # Normalize unicode characters (NFKD decomposes accented chars into base + combining marks)
    normalized = unicodedata.normalize("NFKD", text)
    # Encode to ASCII, ignoring characters that can't be converted
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    # Remove/replace invalid filename characters
    ascii_text = re.sub(r'[<>:"/\\|?*]', "_", ascii_text)
    return ascii_text


def get_mail_merge_placeholders(template_path: Path) -> list[str]:
    """Extract mail merge style placeholders «...» from Word template."""
    doc = Document(template_path)
    text = []

    # Get text from paragraphs
    for para in doc.paragraphs:
        text.append(para.text)

    # Get text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)

    full_text = "\n".join(text)

    # Find «placeholder» patterns (mail merge style)
    placeholders = re.findall(r"«([^»]+)»", full_text)
    return list(set(placeholders))


def replace_placeholder_in_paragraph(paragraph, placeholder: str, value: str):
    """Replace a placeholder in a paragraph, handling split runs."""
    full_text = paragraph.text
    search_text = f"«{placeholder}»"

    if search_text not in full_text:
        return False

    # Simple case: placeholder is in a single run
    for run in paragraph.runs:
        if search_text in run.text:
            run.text = run.text.replace(search_text, str(value))
            return True

    # Complex case: placeholder is split across runs
    # Rebuild the paragraph text
    new_text = full_text.replace(search_text, str(value))

    # Clear all runs and set first run to new text
    if paragraph.runs:
        # Keep formatting from first run
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = new_text
        return True

    return False


def replace_placeholders_in_doc(doc: Document, data: dict):
    """Replace all mail merge placeholders in a document."""
    # Replace in paragraphs
    for para in doc.paragraphs:
        for placeholder, column in PLACEHOLDER_TO_COLUMN.items():
            value = data.get(column, "")
            replace_placeholder_in_paragraph(para, placeholder, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, column in PLACEHOLDER_TO_COLUMN.items():
                        value = data.get(column, "")
                        replace_placeholder_in_paragraph(para, placeholder, value)


def inspect_files():
    """Show Excel columns and template placeholders for debugging."""
    print("=" * 60)
    print("EXCEL FILE INSPECTION")
    print("=" * 60)

    df = pd.read_excel(EXCEL_FILE, sheet_name=EXCEL_SHEET)
    print(f"\nFile: {EXCEL_FILE}")
    print(f"Sheet: {EXCEL_SHEET}")
    print(f"Total rows: {len(df)}")
    print(f"\nColumns ({len(df.columns)}):")
    for i, col in enumerate(df.columns, 1):
        print(f"  {i:2}. {col}")

    print(f"\nFirst row sample:")
    for col in df.columns:
        value = df.iloc[0][col]
        val_str = str(value)[:60] + "..." if len(str(value)) > 60 else str(value)
        print(f"  {col}: {val_str}")

    print("\n" + "=" * 60)
    print("TEMPLATE FILE INSPECTION")
    print("=" * 60)

    print(f"\nFile: {TEMPLATE_FILE}")
    placeholders = get_mail_merge_placeholders(TEMPLATE_FILE)

    print(f"\nMail merge placeholders found ({len(placeholders)}):")
    for p in sorted(placeholders):
        print(f"  «{p}»")

    print("\n" + "=" * 60)
    print("PLACEHOLDER MAPPING")
    print("=" * 60)

    for placeholder in sorted(placeholders):
        if placeholder in PLACEHOLDER_TO_COLUMN:
            col = PLACEHOLDER_TO_COLUMN[placeholder]
            if col in df.columns:
                print(f"  ✓ «{placeholder}» -> '{col}'")
            else:
                print(f"  ✗ «{placeholder}» -> '{col}' (COLUMN NOT FOUND!)")
        else:
            print(f"  ✗ «{placeholder}» -> NOT MAPPED")


def find_libreoffice() -> str | None:
    """Find the LibreOffice executable on this system."""
    # Check PATH first
    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    # Common macOS locations
    mac_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/soffice",
    ]
    for path in mac_paths:
        if Path(path).exists():
            return path

    return None


def convert_docx_to_pdf(docx_path: Path, soffice_path: str, output_dir: Path | None = None) -> Path:
    """Convert a single .docx file to PDF using LibreOffice. Returns the PDF path."""
    if output_dir is None:
        output_dir = docx_path.parent
    output_dir.mkdir(exist_ok=True)

    result = subprocess.run(
        [
            soffice_path,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed for {docx_path.name}:\n{result.stderr}"
        )

    pdf_path = output_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(
            f"PDF was not created for {docx_path.name}. "
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )

    return pdf_path


def convert_all_docx_in_dir(source_dir: Path, dest_dir: Path):
    """Convert all .docx files in source_dir to PDF in dest_dir."""
    soffice_path = find_libreoffice()
    if not soffice_path:
        print(
            "ERROR: LibreOffice not found. Install it with:\n"
            "  brew install --cask libreoffice\n"
            "Then re-run this script.",
            file=sys.stderr,
        )
        sys.exit(1)

    docx_files = sorted(source_dir.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {source_dir}")
        return

    print(f"\nConverting {len(docx_files)} .docx files to PDF...")
    print(f"Output PDF directory: {dest_dir}")
    print(f"Using LibreOffice: {soffice_path}")
    converted = 0
    for i, docx_file in enumerate(docx_files, 1):
        try:
            pdf_path = convert_docx_to_pdf(docx_file, soffice_path, dest_dir)
            print(f"  [{i}/{len(docx_files)}] Converted: {pdf_path.name}")
            converted += 1
        except RuntimeError as e:
            print(f"  [{i}/{len(docx_files)}] FAILED: {e}", file=sys.stderr)
    print(f"\nDone! Converted {converted}/{len(docx_files)} files to PDF.")


def generate_documents(dry_run: bool = False, pdf: bool = False):
    """Generate one document per Excel row."""
    df = pd.read_excel(EXCEL_FILE, sheet_name=EXCEL_SHEET)
    OUTPUT_DIR.mkdir(exist_ok=True)

    print(f"Generating {len(df)} documents...")
    print(f"Output directory: {OUTPUT_DIR}")

    if dry_run:
        print("\n[DRY RUN - No files will be created]\n")

    generated_files = []

    for index, row in df.iterrows():
        # Create data dict from row
        data = row.to_dict()

        # Clean up NaN values
        data = {k: (str(v) if pd.notna(v) else "") for k, v in data.items()}

        # Generate filename: [PI Name] [Nominee name] [Nomination Type]
        pi_name = data.get("PI Name (Ucalgary System)", "").strip()
        nominee_name = data.get("Nominee name", "").strip()
        nomination_type = data.get("Nomination Type", "").strip()

        # Build filename with brackets around each component
        parts = []
        if pi_name:
            parts.append(f"[{normalize_filename(pi_name)}]")
        if nominee_name:
            parts.append(f"[{normalize_filename(nominee_name)}]")
        if nomination_type:
            parts.append(f"[{normalize_filename(nomination_type)}]")

        if parts:
            identifier = " ".join(parts)
        else:
            identifier = f"row_{index + 1}"

        output_filename = OUTPUT_DIR / f"{identifier}.docx"

        if dry_run:
            print(f"  Would create: {output_filename}")
            if pdf:
                print(f"  Would convert to: {output_filename.with_suffix('.pdf')}")
        else:
            # Load fresh copy of template for each document
            doc = Document(TEMPLATE_FILE)
            replace_placeholders_in_doc(doc, data)
            doc.save(output_filename)
            print(f"  Created: {output_filename}")
            generated_files.append(output_filename)

    if not dry_run:
        print(f"\nDone! Generated {len(df)} documents in {OUTPUT_DIR}")

        if pdf and generated_files:
            soffice_path = find_libreoffice()
            if not soffice_path:
                print(
                    "\nERROR: LibreOffice not found. Install it with:\n"
                    "  brew install --cask libreoffice\n"
                    "Then re-run with --pdf-only to convert existing files.",
                    file=sys.stderr,
                )
                sys.exit(1)

            OUTPUT_PDF_DIR.mkdir(exist_ok=True)
            print(f"\nConverting {len(generated_files)} documents to PDF...")
            print(f"Output PDF directory: {OUTPUT_PDF_DIR}")
            print(f"Using LibreOffice: {soffice_path}")
            converted = 0
            for i, docx_file in enumerate(generated_files, 1):
                try:
                    pdf_path = convert_docx_to_pdf(
                        docx_file, soffice_path, OUTPUT_PDF_DIR
                    )
                    print(
                        f"  [{i}/{len(generated_files)}] Converted: {pdf_path.name}"
                    )
                    converted += 1
                except RuntimeError as e:
                    print(
                        f"  [{i}/{len(generated_files)}] FAILED: {e}",
                        file=sys.stderr,
                    )
            print(
                f"\nDone! Converted {converted}/{len(generated_files)} files to PDF."
            )


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word documents from Excel data"
    )
    parser.add_argument(
        "--inspect", action="store_true", help="Inspect files without generating"
    )
    parser.add_argument(
        "--dry-run", action="store_true", help="Show what would be generated"
    )
    parser.add_argument(
        "--pdf",
        action="store_true",
        help="Convert generated .docx files to PDF after creation",
    )
    parser.add_argument(
        "--pdf-only",
        action="store_true",
        help="Convert existing .docx files in output dir to PDF (no generation)",
    )
    args = parser.parse_args()

    if args.inspect:
        inspect_files()
    elif args.pdf_only:
        convert_all_docx_in_dir(OUTPUT_DIR, OUTPUT_PDF_DIR)
    elif args.dry_run:
        generate_documents(dry_run=True, pdf=args.pdf)
    else:
        generate_documents(dry_run=False, pdf=args.pdf)


if __name__ == "__main__":
    main()
